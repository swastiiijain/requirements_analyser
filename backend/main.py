"""FastAPI entrypoint for the secure document-aware chatbot backend."""
from __future__ import annotations

from dotenv import load_dotenv
load_dotenv()

from fastapi import Depends, FastAPI, Header, HTTPException, status, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
import io
import pdfplumber

from backend import auth as auth_utils
from backend import chatbot
from backend import document_reader
from backend.schemas import (
    AskRequest, AskResponse, LoginRequest, LoginResponse, ReadDocResponse, 
    SummariseRequest, SummariseResponse, UploadResponse, AutoSuggestionsResponse,
    ExplainRequest, ExplainResponse, NoteRequest, NoteResponse, NotesListResponse,
    CompareRequest, CompareResponse
)

app = FastAPI(title="Secure Document Chatbot")

# Allow local frontend dev server; adjust in production.
app.add_middleware(
    CORSMiddleware,
    # Allow localhost dev server and Chrome/Edge extension pages
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory mapping from JWT token -> last extracted document text.
session_docs: dict[str, str] = {}
doc_sessions: dict[str, list[dict]] = {}
import uuid

# In-memory storage for notes (in production, use a database)
notes_storage: dict[str, list[dict]] = {}
document_metadata: dict[str, dict] = {}


@app.post("/auth", response_model=LoginResponse, tags=["auth"])
def login(data: LoginRequest):
    if not auth_utils.authenticate_user(data.username, data.password):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid credentials")
    token = auth_utils.create_access_token({"sub": data.username})
    return LoginResponse(access_token=token)


def _parse_bearer(authorization: str | None = Header(None)) -> str:
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Missing or invalid Authorization header")
    return authorization.split(" ", 1)[1]


def _validate_token(token: str) -> None:
    try:
        auth_utils.decode_access_token(token)
    except Exception:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Token expired or invalid")


@app.post("/read-doc", response_model=ReadDocResponse, tags=["document"])
def read_doc(token: str = Depends(_parse_bearer)):
    _validate_token(token)
    filename, text = document_reader.get_active_document_text()
    session_docs[token] = text
    return ReadDocResponse(filename=filename, characters=len(text))


@app.post("/summarise", response_model=SummariseResponse, tags=["document"])
def summarise_page(req: SummariseRequest):
    text = req.text.strip()
    if not text:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No text provided")
    summary = chatbot.summarise(text)
    doc_id = str(uuid.uuid4())
    session_docs[doc_id] = text
    doc_sessions[doc_id] = []  # empty chat history
    return SummariseResponse(summary=summary, characters=len(text), document_id=doc_id)


@app.post("/ask", response_model=AskResponse, tags=["chat"])
def ask_question(req: AskRequest):
    context = session_docs.get(req.document_id)
    if not context:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No document has been analyzed for this session. Click 'Summarise Page' first.")

    answer = chatbot.ask(req.question, context)
    # append to chat log
    doc_sessions.setdefault(req.document_id, []).append({"user": req.question, "bot": answer})
    return AskResponse(answer=answer)


@app.post("/upload", response_model=UploadResponse, tags=["document"])
async def upload_document(file: UploadFile = File(...)):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Only PDF files are supported")

    try:
        # Read the uploaded file content
        content = await file.read()

        # Extract text from PDF
        import pdfplumber
        import io

        text = ""
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text() or ""
                text += extracted + "\n"

        if not text.strip():
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No text found in PDF")

        # Generate summary
        summary = chatbot.summarise(text)

        doc_id = str(uuid.uuid4())
        session_docs[doc_id] = text
        doc_sessions[doc_id] = []
        return UploadResponse(summary=summary, characters=len(text), document_id=doc_id)

    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to process PDF: {str(e)}") 


@app.post("/auto-suggestions", response_model=AutoSuggestionsResponse, tags=["chat"])
def get_auto_suggestions(document_id: str = "default"):
    context = session_docs.get(document_id)
    if not context:
        return AutoSuggestionsResponse(suggestions=["Upload a document first to get suggestions"])
    
    suggestions = chatbot.generate_auto_suggestions(context)
    return AutoSuggestionsResponse(suggestions=suggestions)


@app.post("/explain", response_model=ExplainResponse, tags=["chat"])
def explain_selection(req: ExplainRequest):
    context = session_docs.get(req.document_id or "last", "")
    explanation = chatbot.explain_text(req.text, context)
    return ExplainResponse(explanation=explanation)


@app.post("/notes", response_model=NoteResponse, tags=["notes"])
def save_note(req: NoteRequest):
    import uuid
    from datetime import datetime
    
    note_id = str(uuid.uuid4())
    timestamp = datetime.now().isoformat()
    
    note = {
        "note_id": note_id,
        "content": req.content,
        "document_id": req.document_id,
        "topic": req.topic or "General",
        "timestamp": timestamp
    }
    
    if req.document_id not in notes_storage:
        notes_storage[req.document_id] = []
    
    notes_storage[req.document_id].append(note)
    
    return NoteResponse(**note)


@app.get("/notes", response_model=NotesListResponse, tags=["notes"])
def get_notes(document_id: str = None):
    if document_id:
        notes = notes_storage.get(document_id, [])
    else:
        # Return all notes (fallback when no metadata order)
        notes = []
        if document_metadata:
            for doc_id in sorted(document_metadata.keys()):
                notes.extend(notes_storage.get(doc_id, []))
        else:
            for lst in notes_storage.values():
                notes.extend(lst)
    
    return NotesListResponse(notes=[NoteResponse(**note) for note in notes])


@app.delete("/notes/{note_id}", status_code=204, tags=["notes"])
def delete_note(note_id: str):
    for doc_id, lst in notes_storage.items():
        before = len(lst)
        notes_storage[doc_id] = [n for n in lst if n['note_id'] != note_id]
        if len(notes_storage[doc_id]) < before:
            return
    raise HTTPException(status_code=404, detail="Note not found")


@app.post("/compare", response_model=CompareResponse, tags=["document"])
async def compare_documents_endpoint(
    document1: UploadFile = File(...),
    document2: UploadFile = File(...)
):
    try:
        # Read and extract text from both documents
        content1 = await document1.read()
        content2 = await document2.read()
        
        # Extract text from document 1
        if document1.filename.endswith('.pdf'):
            with io.BytesIO(content1) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    text1 = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif document1.filename.endswith(('.docx', '.doc')):
            from docx import Document
            with io.BytesIO(content1) as docx_buffer:
                doc = Document(docx_buffer)
                text1 = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif document1.filename.endswith('.txt'):
            text1 = content1.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type for document 1: {document1.filename}")
            
        # Extract text from document 2
        if document2.filename.endswith('.pdf'):
            with io.BytesIO(content2) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    text2 = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif document2.filename.endswith(('.docx', '.doc')):
            from docx import Document
            with io.BytesIO(content2) as docx_buffer:
                doc = Document(docx_buffer)
                text2 = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif document2.filename.endswith('.txt'):
            text2 = content2.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type for document 2: {document2.filename}")
            
        if not text1.strip():
            raise HTTPException(status_code=400, detail="Document 1 appears to be empty or unreadable")
        if not text2.strip():
            raise HTTPException(status_code=400, detail="Document 2 appears to be empty or unreadable")
        
        # Compare documents
        summary, changes = chatbot.compare_documents(
            text1, text2, document1.filename, document2.filename
        )
        
        import uuid
        doc1_id = str(uuid.uuid4())
        doc2_id = str(uuid.uuid4())
        
        # Store comparison for later reference
        session_docs[f"compare_{doc1_id}"] = text1
        session_docs[f"compare_{doc2_id}"] = text2
        
        return CompareResponse(
            comparison_summary=summary,
            changes=changes,
            document1_content=text1,
            document2_content=text2,
            document1_id=doc1_id,
            document2_id=doc2_id
        )
    except Exception as e:
        print(f"Compare error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to compare documents: {str(e)}") 